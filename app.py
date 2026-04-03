# ── FIX: Streamlit Cloud uses old SQLite; swap it before chromadb loads ───────
try:
    __import__("pysqlite3")
    import sys
    sys.modules["sqlite3"] = sys.modules.pop("pysqlite3")
except ImportError:
    pass  # Local Windows dev — not needed here

import os
import io
import re
import tempfile
import traceback
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

import streamlit as st

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Research & Strategy Assistant",
    page_icon=None,
    layout="wide",
)

# ── Design tokens from PNG ────────────────────────────────────────────────────
TEAL       = "#1A7A72"
YELLOW     = "#F5E642"
SAGE_GREEN = "#7FB5A8"
CREAM      = "#FFF9C4"   # bright light yellow — matches PNG background exactly
DARK_TEXT  = "#1A3A38"

GLOBAL_CSS = f"""
<style>
/* ── Global background ───────────────────────────────────── */
html, body,
[data-testid="stAppViewContainer"],
[data-testid="stMain"],
[data-testid="stAppViewBlockContainer"] {{
    background-color: {CREAM} !important;
}}
[data-testid="stHeader"] {{
    background-color: {CREAM} !important;
    border-bottom: none !important;
}}
[data-testid="stSidebar"] {{ display: none !important; }}
#MainMenu, footer, header {{ visibility: hidden; }}

.block-container {{
    padding-top: 2rem !important;
    padding-bottom: 6rem !important;
    max-width: 1140px !important;
}}

/* ── Hero banner ─────────────────────────────────────────── */
.hero-banner {{
    background-color: {TEAL};
    border-radius: 22px;
    padding: 32px 48px;
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 28px;
    margin: 0 auto 2.8rem auto;
    width: 100%;
    box-sizing: border-box;
    text-align: center;
}}
.hero-icon {{
    width: 80px;
    height: 80px;
    flex-shrink: 0;
}}
.hero-title {{
    font-family: 'Arial Black', 'Arial', sans-serif;
    font-size: 2.6rem;
    font-weight: 900;
    color: {YELLOW};
    line-height: 1.15;
    margin: 0;
    letter-spacing: -0.5px;
    text-align: left;
}}

/* ── Section headings ────────────────────────────────────── */
.section-heading {{
    font-family: 'Arial Black', 'Arial', sans-serif;
    font-size: 1.55rem;
    font-weight: 900;
    color: {TEAL};
    margin: 0 0 4px 0;
}}
.section-subtext {{
    font-size: 0.88rem;
    color: #444;
    margin-bottom: 10px;
}}

/* ── File uploader ───────────────────────────────────────── */
[data-testid="stFileUploader"] > div {{
    background-color: {SAGE_GREEN} !important;
    border: none !important;
    border-radius: 28px !important;
    padding: 22px 20px !important;
}}
[data-testid="stFileUploaderDropzone"] {{
    background-color: {SAGE_GREEN} !important;
    border-radius: 28px !important;
    border: none !important;
}}
[data-testid="stFileUploaderDropzoneInstructions"] {{
    color: #111111 !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
}}
[data-testid="stFileUploaderDropzoneInstructions"] span,
[data-testid="stFileUploaderDropzoneInstructions"] small,
[data-testid="stFileUploaderDropzoneInstructions"] p,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] small {{
    color: #111111 !important;
}}
/* Browse File button */
[data-testid="stFileUploaderDropzoneInput"] + div button,
[data-testid="stFileUploader"] button {{
    background-color: {YELLOW} !important;
    color: {DARK_TEXT} !important;
    border: none !important;
    border-radius: 30px !important;
    font-weight: 700 !important;
    font-size: 1rem !important;
    padding: 10px 30px !important;
}}

/* ── Text area ───────────────────────────────────────────── */
textarea {{
    background-color: {SAGE_GREEN} !important;
    border: none !important;
    border-radius: 20px !important;
    font-size: 0.95rem !important;
    color: {DARK_TEXT} !important;
    padding: 18px !important;
    box-shadow: none !important;
}}
textarea:focus {{
    box-shadow: none !important;
    border: none !important;
}}

/* ── Primary button (Analyze Business) ───────────────────── */
[data-testid="stBaseButton-primary"] button,
button[kind="primary"],
.stButton > button[kind="primary"],
div[data-testid="stBaseButton-primary"] > button {{
    background-color: {TEAL} !important;
    background: {TEAL} !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 40px !important;
    font-size: 1.15rem !important;
    font-weight: 700 !important;
    padding: 16px 48px !important;
    letter-spacing: 0.2px !important;
    width: 100% !important;
}}
[data-testid="stBaseButton-primary"] button:hover,
button[kind="primary"]:hover {{
    background-color: #145f58 !important;
    background: #145f58 !important;
}}

/* ── Secondary / back button ─────────────────────────────── */
[data-testid="stBaseButton-secondary"] button {{
    background-color: transparent !important;
    border: 2px solid {TEAL} !important;
    color: {TEAL} !important;
    border-radius: 30px !important;
    font-weight: 600 !important;
}}

/* ── Download buttons ────────────────────────────────────── */
[data-testid="stDownloadButton"] button {{
    background-color: {TEAL} !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    padding: 10px 20px !important;
}}
[data-testid="stDownloadButton"] button:hover {{
    background-color: #145f58 !important;
}}

/* ── Footer bar ──────────────────────────────────────────── */
.footer-bar {{
    background-color: {TEAL};
    height: 52px;
    position: fixed;
    bottom: 0;
    left: 0;
    right: 0;
    z-index: 999;
}}

/* ── Results page headings ───────────────────────────────── */
.res-section-title {{
    font-family: 'Arial Black', 'Arial', sans-serif;
    font-size: 1.25rem;
    font-weight: 900;
    color: {TEAL};
    border-left: 5px solid {YELLOW};
    padding-left: 12px;
    margin: 1.6rem 0 0.8rem 0;
}}

/* ── Strategy cards ──────────────────────────────────────── */
[data-testid="stVerticalBlockBorderWrapper"] {{
    border: 2px solid #b2dbd8 !important;
    border-radius: 12px !important;
    background-color: #f0faf9 !important;
}}

/* ── Tabs ────────────────────────────────────────────────── */
[data-testid="stTabs"] [role="tab"] {{
    font-weight: 700 !important;
    color: {TEAL} !important;
}}
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {{
    border-bottom: 3px solid {TEAL} !important;
    color: {TEAL} !important;
}}

/* ── Results header banner ───────────────────────────────── */
.results-banner {{
    background-color: {TEAL};
    border-radius: 14px;
    padding: 16px 28px;
    margin-bottom: 1.5rem;
    display: flex;
    align-items: center;
    gap: 16px;
}}
.results-banner-title {{
    color: {YELLOW};
    font-size: 1.55rem;
    font-weight: 900;
    font-family: 'Arial Black', 'Arial', sans-serif;
    margin: 0;
}}
.results-banner-sub {{
    color: #cceae8;
    font-size: 0.92rem;
    margin-left: auto;
}}
/* ── Scoped: results page body text (dark, readable) ───── */
.results-page p,
.results-page li,
.results-page span,
.results-page div[data-testid="stMarkdownContainer"] p,
.results-page div[data-testid="stMarkdownContainer"] li,
.results-page div[data-testid="stMarkdownContainer"] span,
.results-page div[data-testid="stMarkdownContainer"] ul,
.results-page div[data-testid="stMarkdownContainer"] ol,
.results-page .stMarkdown p,
.results-page .stMarkdown li,
.results-page div[data-testid="stCaptionContainer"] span {{
    color: #1a1a1a !important;
}}

/* Force all standard markdown text (including in tabs) to be dark */
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] span,
[data-testid="stMarkdownContainer"] ul,
[data-testid="stMarkdownContainer"] ol {{
    color: #1a1a1a !important;
}}

/* Scoped: preserve heading colors inside results page */
.results-page .res-section-title {{
    color: {TEAL} !important;
}}
.results-page .results-banner-title {{
    color: {YELLOW} !important;
}}
.results-page .results-banner-sub {{
    color: #cceae8 !important;
}}

/* Scoped: dataframe / table light background */
.results-page [data-testid="stDataFrame"],
.results-page [data-testid="stDataFrame"] > div,
.results-page [data-testid="stDataFrame"] iframe {{
    background-color: #ffffff !important;
    border-radius: 10px !important;
}}
.results-page [data-testid="stDataFrame"] [data-testid="glideDataEditor"],
.results-page [data-testid="stDataFrame"] .dvn-scroller,
.results-page [data-testid="stDataFrame"] canvas {{
    background-color: #ffffff !important;
}}
.results-page [data-testid="stTable"],
.results-page [data-testid="stTable"] table,
.results-page [data-testid="stTable"] th,
.results-page [data-testid="stTable"] td {{
    background-color: #ffffff !important;
    color: #1a1a1a !important;
}}
.results-page .stDataFrame,
.results-page .stDataFrame div {{
    background-color: #ffffff !important;
}}

/* Scoped: download button text stays white */
.results-page [data-testid="stDownloadButton"] button {{
    color: #ffffff !important;
}}

/* ── Spinner text ────────────────────────────────────────── */
[data-testid="stSpinner"] {{
    color: #1A7A72 !important;
    font-weight: bold !important;
}}
[data-testid="stSpinner"] > div > div > div {{
    color: #1A7A72 !important;
}}
</style>
"""

st.markdown(GLOBAL_CSS, unsafe_allow_html=True)


# ── Lazy / guarded imports ────────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def load_heavy_deps():
    try:
        from langchain_community.document_loaders import (
            PyPDFLoader,
            TextLoader,
            Docx2txtLoader,
        )
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain_openai import OpenAIEmbeddings
        from langchain_community.vectorstores import Chroma
        from langchain_core.documents import Document
        from crewai import Agent, Task, Crew, Process

        return {
            "PyPDFLoader": PyPDFLoader,
            "TextLoader": TextLoader,
            "Docx2txtLoader": Docx2txtLoader,
            "RecursiveCharacterTextSplitter": RecursiveCharacterTextSplitter,
            "OpenAIEmbeddings": OpenAIEmbeddings,
            "Chroma": Chroma,
            "Document": Document,
            "Agent": Agent,
            "Task": Task,
            "Crew": Crew,
            "Process": Process,
        }
    except ImportError as e:
        return {"error": str(e)}


# ── Session-state defaults ────────────────────────────────────────────────────
def init_state():
    defaults = {
        "page": "input",
        "uploaded_files_data": [],
        "company_text": "",
        "retriever": None,
        "vectorstore": None,
        "results": None,
        "error": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_state()


# ── Helper: build RAG pipeline ────────────────────────────────────────────────
def build_rag(files_data: list, company_text: str, deps: dict):
    PyPDFLoader = deps["PyPDFLoader"]
    TextLoader = deps["TextLoader"]
    Docx2txtLoader = deps["Docx2txtLoader"]
    RecursiveCharacterTextSplitter = deps["RecursiveCharacterTextSplitter"]
    OpenAIEmbeddings = deps["OpenAIEmbeddings"]
    Chroma = deps["Chroma"]
    Document = deps["Document"]

    all_docs = []

    with tempfile.TemporaryDirectory() as tmp_dir:
        for filename, file_bytes in files_data:
            tmp_path = os.path.join(tmp_dir, filename)
            with open(tmp_path, "wb") as f:
                f.write(file_bytes)
            ext = Path(filename).suffix.lower()
            try:
                if ext == ".pdf":
                    loader = PyPDFLoader(tmp_path)
                elif ext == ".txt":
                    loader = TextLoader(tmp_path, encoding="utf-8")
                elif ext in (".docx", ".doc"):
                    loader = Docx2txtLoader(tmp_path)
                else:
                    st.warning(f"Unsupported file type: {filename}. Skipping.")
                    continue
                docs = loader.load()
                all_docs.extend(docs)
            except Exception as e:
                st.warning(f"Could not load {filename}: {e}")

    if company_text.strip():
        all_docs.append(Document(page_content=company_text.strip(), metadata={"source": "user_input"}))

    if not all_docs:
        raise ValueError("No content to process. Please upload files or enter a company description.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=75)
    chunks = splitter.split_documents(all_docs)

    if not chunks:
        raise ValueError("Text splitting produced no chunks. Please check your input.")

    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

    chroma_dir = tempfile.mkdtemp()
    vectorstore = Chroma.from_documents(chunks, embeddings, persist_directory=chroma_dir)
    retriever = vectorstore.as_retriever(search_kwargs={"k": TOP_K})
    st.session_state["vectorstore"] = vectorstore
    st.session_state["retriever"] = retriever
    st.session_state["chroma_dir"] = chroma_dir

    return retriever


# ── Helper: retrieve context ──────────────────────────────────────────────────
def retrieve_context(retriever, query: str) -> str:
    try:
        docs = retriever.invoke(query)
        return "\n\n".join(d.page_content for d in docs)
    except Exception:
        docs = retriever._get_relevant_documents(query)
        return "\n\n".join(d.page_content for d in docs)


# ── Token limit constants ────────────────────────────────────────────────────
MAX_INPUT_CHARS = 2000   # max user input chars — hard stop if exceeded
MAX_CTX_CHARS   = 3500   # max chars per RAG context chunk (~875 tokens)
TOP_K           = 3      # number of RAG chunks retrieved per query
MAX_TOKENS      = 2200   # output token ceiling per agent call


# ── Helper: run CrewAI pipeline ──────────────────────────────────────────────


def run_crew(retriever, company_text: str, deps: dict) -> dict:
    Agent = deps["Agent"]
    Task = deps["Task"]
    Crew = deps["Crew"]
    Process = deps["Process"]

    openai_api_key = os.environ.get("OPENAI_API_KEY")
    if not openai_api_key:
        raise ValueError("OPENAI_API_KEY not found in environment.")
    os.environ["OPENAI_API_KEY"] = openai_api_key
    model = "openai/gpt-4o-mini"

    research_ctx = retrieve_context(retriever, "key business insights and main topics")[:MAX_CTX_CHARS]
    analysis_ctx = retrieve_context(retriever, "business challenges problems gaps weaknesses")[:MAX_CTX_CHARS]
    strategy_ctx = retrieve_context(retriever, "opportunities improvements recommendations growth")[:MAX_CTX_CHARS]

    # Truncate company text if it exceeds the per-request input limit
    if len(company_text) > MAX_INPUT_CHARS:
        company_text = company_text[:MAX_INPUT_CHARS]

    subject = company_text[:300] if company_text.strip() else "the uploaded documents"

    # Shared LLM config — enforces output token ceiling on every call
    # We use ChatOpenAI so liteLLM doesn't complain about the dictionary format
    from langchain_openai import ChatOpenAI
    llm_instance = ChatOpenAI(
        model="gpt-4o-mini",
        max_tokens=MAX_TOKENS,
    )

    research_agent = Agent(
        role="Senior Research Analyst",
        goal=(
            "Extract 3–5 high-value key insights from business documents. "
            "Output as a tight bullet list. Each bullet must be one sentence "
            "(≤ 20 words). Focus on facts, numbers, and strategic signals only."
        ),
        backstory=(
            "You are a sharp research analyst who distils complex documents into "
            "the fewest, most impactful insights. You never pad output. "
            "Bullets only — no headings, no paragraphs."
        ),
        llm=llm_instance,
        verbose=False,
        allow_delegation=False,
    )

    analysis_agent = Agent(
        role="Business Analysis Expert",
        goal=(
            "Identify exactly 5 business pain points from the context. "
            "For each, output: PROBLEM, IMPACT, PRIORITY. "
            "Keep descriptions short (≤ 15 words per field). No extra commentary."
        ),
        backstory=(
            "You are a business analyst who surfaces structured pain points. "
            "You output concise, table-ready blocks. No intro text, no padding."
        ),
        llm=llm_instance,
        verbose=False,
        allow_delegation=False,
    )

    strategy_agent = Agent(
        role="Strategic Consultant",
        goal=(
            "Recommend exactly 5 actionable strategies. "
            "Each strategy: FIX title + 2 action bullets + OUTCOME sentence. "
            "Aim for clarity and brevity. Skip any preamble."
        ),
        backstory=(
            "You are a management consultant who writes crisp strategy cards. "
            "Every card delivers a title, two concrete actions, and one outcome. "
            "No paragraphs, no repetition."
        ),
        llm=llm_instance,
        verbose=False,
        allow_delegation=False,
    )

    content_agent = Agent(
        role="Professional Content Writer",
        goal=(
            "Produce a structured report with exactly five sections: "
            "(1) Summary — short paragraph, "
            "(2) Key Insights — 3–5 bullets, "
            "(3) Pain Points — 5 bullets, "
            "(4) Strategies — 5 bullets, "
            "(5) Email — max 150 words. "
            "Use the exact section markers provided. Stay concise throughout."
        ),
        backstory=(
            "You are a business writer who assembles polished, token-efficient reports. "
            "You follow section markers exactly, avoid repetition, and never exceed "
            "the requested length for any section."
        ),
        llm=llm_instance,
        verbose=False,
        allow_delegation=False,
    )

    research_task = Task(
        description=(
            f"Business context: {subject}\n\n"
            f"DOCUMENT CONTEXT:\n{research_ctx}\n\n"
            "OUTPUT RULES (STRICT):\n"
            "- Output ONLY a bullet list of 5-7 key insights\n"
            "- Each bullet = one short sentence (max 15 words)\n"
            "- Start each line with '- '\n"
            "- NO headings, NO paragraphs, NO explanations\n"
            "- Only facts and insights from the context\n\n"
            "Example format:\n"
            "- Company operates in 3 markets with declining margins in two.\n"
            "- Customer churn rate is above industry average at 18%.\n"
            "- Product roadmap lacks mobile-first features despite 60% mobile users."
        ),
        expected_output="A bullet list of 5-7 concise key insights, one sentence each.",
        agent=research_agent,
    )

    analysis_task = Task(
        description=(
            f"DOCUMENT CONTEXT:\n{analysis_ctx}\n\n"
            "OUTPUT RULES (STRICT):\n"
            "- Identify exactly 5 business pain points\n"
            "- Format EACH pain point as follows (use this exact structure):\n\n"
            "PROBLEM: [short name of the problem]\n"
            "IMPACT: [one sentence on business impact]\n"
            "PRIORITY: [High / Medium / Low]\n"
            "---\n\n"
            "- NO paragraphs\n"
            "- NO extra commentary before or after\n"
            "- Repeat the block exactly 5 times"
        ),
        expected_output="5 structured pain point blocks, each with PROBLEM, IMPACT, and PRIORITY fields.",
        agent=analysis_agent,
        context=[research_task],
    )

    strategy_task = Task(
        description=(
            f"DOCUMENT CONTEXT:\n{strategy_ctx}\n\n"
            "OUTPUT RULES (STRICT):\n"
            "- Provide exactly 5 strategy recommendations\n"
            "- Format EACH strategy as follows:\n\n"
            "FIX: [Short strategy title]\n"
            "- [Action step 1]\n"
            "- [Action step 2]\n"
            "- [Action step 3]\n"
            "OUTCOME: [One sentence on expected result]\n"
            "---\n\n"
            "- NO paragraphs\n"
            "- NO extra text before or after\n"
            "- Repeat the block exactly 5 times"
        ),
        expected_output="5 strategy blocks each with FIX title, 2-3 action bullets, and OUTCOME.",
        agent=strategy_agent,
        context=[analysis_task],
    )

    content_task = Task(
        description=(
            "Using all prior agent outputs, produce a structured report with "
            "EXACTLY FIVE sections. Copy the section markers verbatim.\n\n"
            "===SUMMARY===\n"
            "[One short paragraph — 3–4 sentences capturing the business situation.]\n\n"
            "===KEY INSIGHTS===\n"
            "- [Insight 1]\n"
            "- [Insight 2]\n"
            "- [Insight 3]\n"
            "(3–5 bullets, each ≤ 20 words)\n\n"
            "===PAIN POINTS===\n"
            "- [Pain point 1]\n"
            "- [Pain point 2]\n"
            "- [Pain point 3]\n"
            "- [Pain point 4]\n"
            "- [Pain point 5]\n"
            "(Exactly 5 bullets, each ≤ 20 words)\n\n"
            "===STRATEGIES===\n"
            "- [Strategy 1]\n"
            "- [Strategy 2]\n"
            "- [Strategy 3]\n"
            "- [Strategy 4]\n"
            "- [Strategy 5]\n"
            "(Exactly 5 bullets, each ≤ 20 words)\n\n"
            "===EMAIL===\n"
            "Subject: [Email subject line]\n\n"
            "Dear [Stakeholder],\n\n"
            "[Up to 150 words. Cover top finding, key risk, and one call to action.]\n\n"
            "Best regards,\n"
            "[Your Name]\n\n"
            "TOKEN RULES:\n"
            "- Total output must stay well under 2000 tokens\n"
            "- Use the exact section markers above — no extra headings\n"
            "- No repetition across sections\n"
            "- Skip any section intro text; go straight to content"
        ),
        expected_output=(
            "Five sections: ===SUMMARY===, ===KEY INSIGHTS===, ===PAIN POINTS===, "
            "===STRATEGIES===, ===EMAIL=== — each concise and within token limits."
        ),
        agent=content_agent,
        context=[research_task, analysis_task, strategy_task],
    )

    crew = Crew(
        agents=[research_agent, analysis_agent, strategy_agent, content_agent],
        tasks=[research_task, analysis_task, strategy_task, content_task],
        process=Process.sequential,
        verbose=False,
    )

    crew.kickoff()

    def safe_output(task):
        try:
            return str(task.output.raw) if hasattr(task.output, "raw") else str(task.output)
        except Exception:
            return ""

    return {
        "research": safe_output(research_task),
        "analysis": safe_output(analysis_task),
        "strategy": safe_output(strategy_task),
        "content": safe_output(content_task),
    }


# ── Output parsers ────────────────────────────────────────────────────────────
def parse_pain_points(text: str) -> list:
    blocks = re.split(r"-{3,}", text)
    rows = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        problem = re.search(r"PROBLEM\s*:\s*(.+)", block, re.IGNORECASE)
        impact = re.search(r"IMPACT\s*:\s*(.+)", block, re.IGNORECASE)
        priority = re.search(r"PRIORITY\s*:\s*(.+)", block, re.IGNORECASE)
        if problem or impact or priority:
            rows.append({
                "Problem": problem.group(1).strip() if problem else "-",
                "Impact": impact.group(1).strip() if impact else "-",
                "Priority": priority.group(1).strip() if priority else "-",
            })
    return rows


def parse_strategies(text: str) -> list:
    blocks = re.split(r"-{3,}", text)
    strategies = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        fix_match = re.search(r"FIX\s*:\s*(.+)", block, re.IGNORECASE)
        outcome_match = re.search(r"OUTCOME\s*:\s*(.+)", block, re.IGNORECASE)
        bullets = re.findall(r"^[\-\*]\s+(.+)", block, re.MULTILINE)
        fix_title = fix_match.group(1).strip() if fix_match else ""
        outcome_text = outcome_match.group(1).strip() if outcome_match else ""
        bullets = [b for b in bullets if b != fix_title and b != outcome_text]
        if fix_match or bullets:
            strategies.append({
                "title": fix_title or "Strategy",
                "steps": bullets,
                "outcome": outcome_text,
            })
    return strategies


def parse_content_sections(text: str) -> dict:
    """Extract the five structured sections from the content agent output."""
    def extract(marker_start, marker_end=None):
        if marker_end:
            pattern = rf"==={marker_start}===\s*(.*?)\s*==={marker_end}==="
        else:
            pattern = rf"==={marker_start}===\s*(.*)"
        m = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        return m.group(1).strip() if m else ""

    return {
        "summary":      extract("SUMMARY",      "KEY INSIGHTS"),
        "key_insights": extract("KEY INSIGHTS",  "PAIN POINTS"),
        "pain_points":  extract("PAIN POINTS",   "STRATEGIES"),
        "strategies":   extract("STRATEGIES",    "EMAIL"),
        "email":        extract("EMAIL"),
    }


# ── Download generators ───────────────────────────────────────────────────────
def build_txt(results: dict) -> str:
    return "\n\n".join(filter(None, [
        "AI RESEARCH & STRATEGY ASSISTANT REPORT",
        "=" * 50,
        "KEY INSIGHTS\n" + results.get("research", ""),
        "PAIN POINTS & GAPS\n" + results.get("analysis", ""),
        "STRATEGY SUGGESTIONS\n" + results.get("strategy", ""),
        "GENERATED OUTPUT\n" + results.get("content", ""),
    ]))


def build_pdf_bytes(results: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2.5*cm, bottomMargin=2.5*cm,
        )

        styles = getSampleStyleSheet()
        teal_color = colors.HexColor("#1A7A72")

        title_style = ParagraphStyle("RT", parent=styles["Title"],
                                     fontSize=20, textColor=teal_color, spaceAfter=6)
        h1_style = ParagraphStyle("H1", parent=styles["Heading1"],
                                  fontSize=14, textColor=teal_color,
                                  spaceBefore=14, spaceAfter=4)
        body_style = ParagraphStyle("B", parent=styles["Normal"],
                                    fontSize=10, leading=15, spaceAfter=4)

        story = [
            Paragraph("AI Research &amp; Strategy Assistant", title_style),
            Paragraph("Business Analysis Report", styles["Normal"]),
            Spacer(1, 0.4*cm),
            HRFlowable(width="100%", thickness=2, color=teal_color),
            Spacer(1, 0.4*cm),
        ]

        sections = [
            ("Key Insights", results.get("research", "")),
            ("Pain Points & Gaps", results.get("analysis", "")),
            ("Strategy Suggestions", results.get("strategy", "")),
            ("Generated Output", results.get("content", "")),
        ]

        for heading, content in sections:
            if content.strip():
                story.append(Paragraph(heading, h1_style))
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=colors.HexColor("#cccccc")))
                story.append(Spacer(1, 0.2*cm))
                for line in content.splitlines():
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 0.15*cm))
                        continue
                    safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if line.startswith("## "):
                        story.append(Paragraph(safe[3:], h1_style))
                    elif line.startswith("# "):
                        story.append(Paragraph(safe[2:], h1_style))
                    elif line.startswith("- ") or line.startswith("* "):
                        story.append(Paragraph(f"&bull; {safe[2:]}", body_style))
                    else:
                        story.append(Paragraph(safe, body_style))
                story.append(Spacer(1, 0.3*cm))

        doc.build(story)
        return buf.getvalue()
    except ImportError:
        return b""


def build_docx_bytes(results: dict) -> bytes:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        teal_rgb = RGBColor(0x1A, 0x7A, 0x72)

        title_para = doc.add_paragraph()
        title_run = title_para.add_run("AI Research & Strategy Assistant")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = teal_rgb
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        sub_para = doc.add_paragraph("Business Analysis Report")
        sub_para.runs[0].font.size = Pt(11)
        doc.add_paragraph()

        sections = [
            ("Key Insights", results.get("research", "")),
            ("Pain Points & Gaps", results.get("analysis", "")),
            ("Strategy Suggestions", results.get("strategy", "")),
            ("Generated Output", results.get("content", "")),
        ]

        for heading, content in sections:
            if content.strip():
                h = doc.add_heading(heading, level=1)
                for run in h.runs:
                    run.font.color.rgb = teal_rgb

                for line in content.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("## ") or line.startswith("# "):
                        sub_h = doc.add_heading(line.lstrip("# "), level=2)
                        for run in sub_h.runs:
                            run.font.color.rgb = teal_rgb
                    elif line.startswith("- ") or line.startswith("* "):
                        p = doc.add_paragraph(line[2:], style="List Bullet")
                        if p.runs:
                            p.runs[0].font.size = Pt(10)
                    else:
                        p = doc.add_paragraph(line)
                        if p.runs:
                            p.runs[0].font.size = Pt(10)
                doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return b""


# ── Hero SVG icon ─────────────────────────────────────────────────────────────
HERO_ICON_SVG = """
<svg viewBox="0 0 80 80" fill="none" xmlns="http://www.w3.org/2000/svg" class="hero-icon">
  <!-- Gear outer -->
  <path d="M40 8 L44 2 L48 8 L55 6 L56 13 L63 14 L62 21 L68 25 L65 31
           L70 37 L66 42 L69 49 L63 52 L63 59 L56 60 L53 67 L46 65
           L40 70 L34 65 L27 67 L24 60 L17 59 L17 52 L11 49 L14 42
           L10 37 L15 31 L12 25 L18 21 L17 14 L24 13 L25 6 L32 8 Z"
        fill="none" stroke="{YELLOW}" stroke-width="3" stroke-linejoin="round"/>
  <!-- Inner circle -->
  <circle cx="40" cy="37" r="14" fill="none" stroke="{YELLOW}" stroke-width="3"/>
  <!-- Lightbulb body -->
  <path d="M34 34 Q34 27 40 27 Q46 27 46 34 Q46 39 43 41 L43 45 L37 45 L37 41 Q34 39 34 34Z"
        fill="{YELLOW}"/>
  <!-- Lightbulb base lines -->
  <rect x="37" y="45" width="6" height="2" rx="1" fill="{YELLOW}"/>
  <rect x="37.5" y="47.5" width="5" height="2" rx="1" fill="{YELLOW}"/>
</svg>
""".replace("{YELLOW}", YELLOW)


# ══════════════════════════════════════════════════════════════════════════════
# INPUT PAGE
# ══════════════════════════════════════════════════════════════════════════════
def page_input():
    # Hero banner — full width
    _hero_html = (
        '<div class="hero-banner">'
        + HERO_ICON_SVG
        + '<div class="hero-title">AI Research &amp; Strategy Assistant</div>'
        + '</div>'
    )
    st.markdown(_hero_html, unsafe_allow_html=True)

    openai_api_key = os.environ.get("OPENAI_API_KEY")
    if not openai_api_key:
        st.error("OPENAI_API_KEY not found in environment. Please set it in your .env file.")
        return

    # Two-column layout — matching PNG exactly
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<p class="section-heading">Upload Documents</p>', unsafe_allow_html=True)
        st.markdown('<p class="section-subtext">PDF, TXT, or DOCX</p>', unsafe_allow_html=True)
        uploaded = st.file_uploader(
            "Drag and drop files here  •  Limit 200MB per file  •  PDF, TXT, DOCX, DOC",
            type=["pdf", "txt", "docx", "doc"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )

    with col2:
        st.markdown('<p class="section-heading">Company Description</p>', unsafe_allow_html=True)
        st.markdown('<p class="section-subtext">Describe the company, products, market, and challenges:</p>', unsafe_allow_html=True)
        company_text = st.text_area(
            "company_description",
            height=185,
            placeholder="",
            label_visibility="collapsed",
        )

    st.markdown("<div style='margin-top:2.2rem;'></div>", unsafe_allow_html=True)

    # Full-width Analyze Business button (matches PNG)
    run_btn = st.button("Analyze Business", type="primary", use_container_width=True)

    if run_btn:
        if not uploaded and not company_text.strip():
            st.error("Please upload at least one document or enter a company description.")
            return

        # ── Hard input size safeguard ─────────────────────────────────────────
        if len(company_text) > MAX_INPUT_CHARS:
            st.warning(
                f"⚠️ Company description exceeds {MAX_INPUT_CHARS} characters "
                f"({len(company_text)} entered). Please shorten it and try again."
            )
            return

        with st.spinner("🔍 Analyzing your business — this may take 1–2 minutes..."):
            st.session_state["uploaded_files_data"] = [(f.name, f.read()) for f in (uploaded or [])]
            st.session_state["company_text"] = company_text
            st.session_state["results"] = None
            st.session_state["error"] = None

            deps = load_heavy_deps()
            if "error" in deps:
                st.error(f"Dependency import failed: {deps['error']}")
                st.info(
                    "Run: pip install streamlit langchain langchain-community "
                    "crewai chromadb pypdf docx2txt "
                    "reportlab python-docx openai langchain-openai python-dotenv langchain-text-splitters"
                )
                return

            with st.spinner("Building RAG pipeline (embeddings + vector store)..."):
                try:
                    retriever = build_rag(
                        st.session_state["uploaded_files_data"],
                        company_text,
                        deps,
                    )
                    st.session_state["retriever"] = retriever
                except Exception as e:
                    st.error(f"RAG pipeline error: {e}\n\n{traceback.format_exc()}")
                    return

            with st.spinner("Running multi-agent analysis (Research > Analysis > Strategy > Content)..."):
                try:
                    results = run_crew(
                        st.session_state["retriever"],
                        company_text,
                        deps,
                    )
                    st.session_state["results"] = results
                except Exception as e:
                    st.error(f"CrewAI pipeline error: {e}\n\n{traceback.format_exc()}")
                    return

            st.session_state["page"] = "results"
            st.rerun()

    st.markdown('<div class="footer-bar"></div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# RESULTS PAGE
# ══════════════════════════════════════════════════════════════════════════════
def page_results():
    st.markdown('<div class="results-page">', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="results-banner">
            <span class="results-banner-title">AI Research &amp; Strategy Assistant</span>
            <span class="results-banner-sub">Analysis Results</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.button("Back to Input"):
        st.session_state["page"] = "input"
        st.rerun()

    results = st.session_state.get("results")
    if not results:
        st.warning("No results found. Please go back and run the analysis.")
        return

    st.markdown("<div style='margin-top:0.5rem;'></div>", unsafe_allow_html=True)

    # ── 1. Key Insights ───────────────────────────────────────────────────────
    st.markdown('<p class="res-section-title">Key Insights</p>', unsafe_allow_html=True)
    research_out = results.get("research", "").strip()
    if research_out:
        lines = [l.strip() for l in research_out.splitlines() if l.strip()]
        bullet_lines = [l for l in lines if l.startswith("-") or l.startswith("*") or l.startswith("•")]
        display_lines = bullet_lines if bullet_lines else lines
        for line in display_lines[:7]:
            clean = re.sub(r"^[-*•]\s*", "", line)
            st.markdown(f"- {clean}")
    else:
        st.info("No research output available.")

    st.divider()

    # ── 2. Pain Points Table ──────────────────────────────────────────────────
    st.markdown('<p class="res-section-title">Pain Points & Gaps</p>', unsafe_allow_html=True)
    analysis_out = results.get("analysis", "").strip()
    if analysis_out:
        rows = parse_pain_points(analysis_out)
        if rows:
            import pandas as pd
            df = pd.DataFrame(rows)

            def colour_priority(val):
                colours = {"High": "#ff4b4b", "Medium": "#ffa500", "Low": "#21c354"}
                c = colours.get(val.strip(), "#888888")
                return f"color: {c}; font-weight: bold;"

            styled = df.style.applymap(colour_priority, subset=["Priority"])
            st.dataframe(styled, use_container_width=True, hide_index=True)
        else:
            st.markdown(analysis_out)
    else:
        st.info("No analysis output available.")

    st.divider()

    # ── 3. Strategy Recommendations ──────────────────────────────────────────
    st.markdown('<p class="res-section-title">Strategy Recommendations</p>', unsafe_allow_html=True)
    strategy_out = results.get("strategy", "").strip()
    if strategy_out:
        strategies = parse_strategies(strategy_out)
        if strategies:
            num_cols = min(len(strategies), 3)
            cols = st.columns(num_cols)
            for i, strat in enumerate(strategies[:5]):
                with cols[i % num_cols]:
                    with st.container(border=True):
                        st.markdown(f"**{strat['title']}**")
                        for step in strat["steps"]:
                            st.markdown(f"- {step}")
                        if strat["outcome"]:
                            st.caption(f"Outcome: {strat['outcome']}")
        else:
            st.markdown(strategy_out)
    else:
        st.info("No strategy output available.")

    st.divider()

    # ── 4. Generated Output (Tabs) ────────────────────────────────────────────
    st.markdown('<p class="res-section-title">Generated Output</p>', unsafe_allow_html=True)
    content_out = results.get("content", "").strip()

    if content_out:
        sections = parse_content_sections(content_out)
        has_parsed = any(sections.values())

        tab_summary, tab_insights, tab_email = st.tabs(["Summary", "Key Insights & Pain Points", "Email"])

        with tab_summary:
            summary_text = sections.get("summary", "") if has_parsed else ""
            if summary_text:
                st.markdown(summary_text)
            else:
                st.markdown(content_out)

        with tab_insights:
            ki_text = sections.get("key_insights", "") if has_parsed else ""
            pp_text = sections.get("pain_points", "") if has_parsed else ""
            strat_text = sections.get("strategies", "") if has_parsed else ""
            if ki_text:
                st.markdown("**Key Insights**")
                st.markdown(ki_text)
            if pp_text:
                st.markdown("**Pain Points**")
                st.markdown(pp_text)
            if strat_text:
                st.markdown("**Strategies**")
                st.markdown(strat_text)
            if not ki_text and not pp_text and not strat_text:
                st.markdown(content_out)

        with tab_email:
            email_text = sections.get("email", "") if has_parsed else ""
            if email_text:
                st.markdown(email_text)
            else:
                st.markdown(content_out)
    else:
        st.info("No content output available.")

    st.divider()

    # ── Download ──────────────────────────────────────────────────────────────
    st.markdown('<p class="res-section-title">Download Report</p>', unsafe_allow_html=True)

    txt_data  = build_txt(results)
    pdf_data  = build_pdf_bytes(results)
    docx_data = build_docx_bytes(results)

    dl1, dl2, dl3 = st.columns(3)

    with dl1:
        st.download_button(
            label="Download as TXT",
            data=txt_data,
            file_name="ai_strategy_report.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with dl2:
        if pdf_data:
            st.download_button(
                label="Download as PDF",
                data=pdf_data,
                file_name="ai_strategy_report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        else:
            st.warning("PDF unavailable — run: pip install reportlab")
    with dl3:
        if docx_data:
            st.download_button(
                label="Download as Word",
                data=docx_data,
                file_name="ai_strategy_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.warning("Word unavailable — run: pip install python-docx")

    st.markdown('</div>', unsafe_allow_html=True)  # close .results-page
    st.markdown('<div class="footer-bar"></div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# ROUTER
# ══════════════════════════════════════════════════════════════════════════════
def main():
    page = st.session_state.get("page", "input")
    if page == "results" and st.session_state.get("results"):
        page_results()
    else:
        page_input()


if __name__ == "__main__":
    main()
